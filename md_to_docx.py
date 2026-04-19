from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# Estilos base
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

def set_heading(paragraph, level, text):
    sizes = {1: 20, 2: 16, 3: 13}
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes.get(level, 12))
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    else:
        run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)

def add_table_from_lines(doc, lines):
    rows = [l.strip() for l in lines if l.strip().startswith('|')]
    if len(rows) < 2:
        return
    # quitar separadores |---|---|
    rows = [r for r in rows if not re.match(r'^\|[-|\s:]+\|$', r)]
    parsed = []
    for r in rows:
        cells = [c.strip() for c in r.strip('|').split('|')]
        parsed.append(cells)
    if not parsed:
        return
    cols = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=cols)
    table.style = 'Table Grid'
    for i, row in enumerate(parsed):
        for j, cell_text in enumerate(row):
            if j < cols:
                cell = table.rows[i].cells[j]
                cell.text = cell_text
                if i == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
    doc.add_paragraph()

def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run('\n'.join(lines))
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    # fondo gris claro via shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)

md_path = '/Users/santiagobruschera/Desktop/santi facultad/quinto semestre/BDD3/tarea 2/lab2-bdd/INFORME.md'
with open(md_path, encoding='utf-8') as f:
    raw_lines = f.readlines()

lines = [l.rstrip('\n') for l in raw_lines]

i = 0
while i < len(lines):
    line = lines[i]

    # Bloque de código
    if line.strip().startswith('```'):
        code_lines = []
        i += 1
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1
        add_code_block(doc, code_lines)
        i += 1
        continue

    # Tabla
    if line.strip().startswith('|'):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i])
            i += 1
        add_table_from_lines(doc, table_lines)
        continue

    # Encabezados
    m = re.match(r'^(#{1,3})\s+(.*)', line)
    if m:
        level = len(m.group(1))
        text = m.group(2)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        set_heading(p, level, text)
        i += 1
        continue

    # Separador ---
    if re.match(r'^---+$', line.strip()):
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'AAAAAA')
        pBdr.append(bottom)
        pPr.append(pBdr)
        i += 1
        continue

    # Línea vacía
    if not line.strip():
        doc.add_paragraph()
        i += 1
        continue

    # Bullet list
    if re.match(r'^[-*]\s+', line.strip()):
        text = re.sub(r'^[-*]\s+', '', line.strip())
        p = doc.add_paragraph(style='List Bullet')
        # inline bold/code
        parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            else:
                p.add_run(part)
        i += 1
        continue

    # Párrafo normal con inline formatting
    p = doc.add_paragraph()
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', line)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            p.add_run(part)
    i += 1

out_path = '/Users/santiagobruschera/Desktop/santi facultad/quinto semestre/BDD3/tarea 2/lab2-bdd/INFORME.docx'
doc.save(out_path)
print(f'Guardado en: {out_path}')
